# this code finds sentences with user provided keywords
import re
import string
from termcolor import colored
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def prRed(skk):
  for wds in matches:
    skk = skk.replace(wds,"\033[44;33m{}\033[00m".format(wds))
  return skk
def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )
# Create a regex pattern to match a string
# that starts with one or more special characters
pattern = r'^[' + string.punctuation + ']+'

paper_dir = "all_papers.txt"

all_papers_file = open(paper_dir,"r",encoding="utf-8")
lines = [line.rstrip() for line in all_papers_file]

#matches = ["filamin","actinin","fimbrin", "fascin"]
#matches = ["contractile ring"]
matches = ["bind to"] # add your keywords here
paper_content = [''.join(lines)]
sentences = paper_content[0].split(".") #sentences is a list of strings splitted by .
i = 0

document = Document()
document.add_heading('sentences with keywords', level=1)


for item in sentences: #item is string
  if all(x in item for x in matches): # replace alll with any or vv
    each_line = item.split(" ")
    while("" in each_line):
      each_line.remove("")
# Check if string starts with special characters
    if re.search(pattern, " ".join(each_line)) is None:
      i+=1
      sent = " ".join(each_line)
      cleaned_sent = ''.join(c for c in sent if valid_xml_char_ordinal(c))
      sent = sent.replace(matches[0], prRed(sent))
      print(i,". ",sent)
      print(" ")      
      document.add_paragraph(str(i)+". "+cleaned_sent)

all_papers_file.close()

document.save('docx_file.docx')
